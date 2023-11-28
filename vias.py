from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tkinter as tk
from tkinter import ttk, filedialog
import os

def criar_documento(file_name, titulo, assinatura1, assinatura2):
    document = Document()

    # Adding the image
    img_path = 'brasao.png'  # Replace with the path to your image file
    document.add_picture(img_path, width=Inches(1.0))  # Adjust width as needed
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Adding title
    title = document.add_paragraph()

    title.alignment = 1  # Center alignment

    title_run_1 = title.add_run(f"""Tribunal Regional Eleitoral do Rio Grande do Norte - TRE/RN
Secretaria de Administração
Coordenadoria de Material e Patrimônio
Seção de Patrimônio""")

    title_run_1.font.size = Pt(12)

    title_run_2 = title.add_run(f"\n\n{titulo}\n")
    title_run_2.font.size = Pt(14)
    title_run_2.bold = True

    entry_values = [entry.get() for entry in entries]

    # Adding sections for ORIGEM and DESTINO
    document.add_paragraph(f"ORIGEM: {entry_values[0]}")
    document.add_paragraph(f"DESTINO: {entry_values[1]}")
    document.add_paragraph()

    # Adding a table
    table = document.add_table(rows=1, cols=3)  # Initial row for headers
    table.style = "Table Grid"  # Apply table style

    # Adding headers to the table
    headers = table.rows[0].cells
    headers[0].text = 'ITEM'
    headers[1].text = 'DESCRIÇÃO'
    headers[2].text = 'QNTD'

    # Setting preferred widths for table columns
    table.autofit = False
    table.columns[0].width = Inches(0.7)  # Adjust width as needed
    table.columns[1].width = Inches(5)  # Adjust width as needed
    table.columns[2].width = Inches(0.7)  # Adjust width as needed

    # Justify text in header cells to the middle
    for header in headers:
        for paragraph in header.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(12)  # Adjust font size if needed
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run.bold = True  # Make the text bold

    # Center the entire table within the document
    table.alignment = 1  # 0=left, 1=center, 2=right

    # Add data from Treeview to the table
    for item in tree.get_children():
        values = tree.item(item, 'values')
        row_cells = table.add_row().cells
        for index, value in enumerate(values):
            row_cells[index].text = value
            p = row_cells[index].paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_paragraph(
        "\n\n\n_____/_____/________                                    _____/_____/________                                    _____/_____/________\n").bold = True
    document.add_paragraph(
        "____________________                                      ____________________                                    ____________________").bold = True
    document.add_paragraph(
        f"   PATRIMÔNIO                                                 {assinatura1}                                                {assinatura2}").bold = True

    # Save the document
    document.save(file_name)


def add_item():
    entry_values = [entry.get() for entry in entries]

    if entry_values[2] and entry_values[3] and entry_values[4]:

        # Insert values into treeview
        tree.insert("", "end", values=(entry_values[2], entry_values[4].upper(), entry_values[3]))


def remove_item():
    selected_item = tree.selection()
    for item in selected_item:
        tree.delete(item)


def edit_item():
    entry_values = [entry.get() for entry in entries]

    selected_item = tree.selection()
    for item in selected_item:
        tree.item(item, values=(entry_values[2], entry_values[4].upper(), entry_values[3]))


def gerar_saida():
    criar_documento("saida.docx", "AUTORIZAÇÃO DE SAIDA DE MATERIAL", "PORTADOR", "PORTARIA")


def gerar_transferencia():
    criar_documento("transferencia.docx", "TERMO DE TRANSFERÊNCIA INTERNA", "CEDENTE", "RECEBEDOR")


def center_window(window, min_width, min_height):
    # Get the screen width and height
    screen_width = window.winfo_screenwidth()
    screen_height = window.winfo_screenheight()

    # Calculate the x and y coordinates for centering
    x = (screen_width - min_width) // 2
    y = (screen_height - min_height) // 2

    # Set the window's geometry to center it on the screen
    window.geometry(f"{min_width}x{min_height}+{x}+{y}")


root = tk.Tk()
root.title("Gerador de Vias Avulsas")
root.resizable(False, False)
root.geometry("500x500")

if os.path.isfile("documento.ico"):
    root.iconbitmap("documento.ico")

# Frames for organizing widgets
labels_entries_frame = tk.Frame(root)
labels_entries_frame.pack(fill="x")
labels_entries_frame.grid_columnconfigure(1, weight=1)

buttons_frame = tk.Frame(root)
buttons_frame.pack(fill="x")

# Label texts and respective entries using grid
label_texts = ["Origem:", "Destino:", "Item:", "Quantidade:", "Descrição:"]
entries = []

for i, label_text in enumerate(label_texts):
    label = tk.Label(labels_entries_frame, text=label_text)
    label.grid(row=i, column=0, padx=5, pady=5, sticky="w")

    entry = tk.Entry(labels_entries_frame)
    entry.grid(row=i, column=1, padx=5, pady=5, sticky="ew")
    entries.append(entry)

# Buttons using pack
button_texts = ["Adicionar", "Remover", "Editar"]
commands = [add_item, remove_item, edit_item]

for button_text, command in zip(button_texts, commands):
    button = tk.Button(buttons_frame, text=button_text, command=command)
    button.pack(side="left", padx=5, pady=5)

# Treeview
tree = ttk.Treeview(root, columns=("Item", "Descrição", "QTD"))
tree.heading("Item", text="Item")
tree.heading("Descrição", text="Descrição")
tree.heading("QTD", text="QTD")

# Set different column widths
tree.column("Item", width=50)
tree.column("Descrição", width=350)
tree.column("QTD", width=50)
tree.column("#0", width=1)  # Hide the ID column

tree.pack(fill="both", expand=True, padx=5, pady=5)

# Menu
menu_bar = tk.Menu(root)
root.config(menu=menu_bar)

gerar_menu = tk.Menu(menu_bar, tearoff=0)
menu_bar.add_cascade(label="Menu", menu=gerar_menu)
gerar_menu.add_command(label="Gerar Autorização de Saída", command=gerar_saida)
gerar_menu.add_command(label="Gerar Transferência Interna", command=gerar_transferencia)

center_window(root, 500, 500)

root.mainloop()
