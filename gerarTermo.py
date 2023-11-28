from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tkinter import messagebox, filedialog
import os


def criar_documento(titulo, assinatura1, assinatura2, entries, tree):
    document = Document()

    # Adding a table with 1 row and 2 columns
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Setting preferred widths for table columns
    table.columns[0].width = Inches(1.3)  # Adjust width as needed
    table.columns[1].width = Inches(4.7)  # Adjust width as needed

    # Access the cells
    cell_image = table.cell(0, 0)
    cell_title = table.cell(0, 1)

    # Adding the image to the first cell
    img_path = 'brasao.png'  # Replace with the path to your image file
    cell_image_paragraph = cell_image.paragraphs[0]
    run = cell_image_paragraph.add_run()
    run.add_picture(img_path, width=Inches(1.0))  # Adjust width as needed

    # Adding the title to the second cell
    title_text = (
        "Tribunal Regional Eleitoral do Rio Grande do Norte - TRE/RN\n"
        "Secretaria de Administração\n"
        "Coordenadoria de Material e Patrimônio\n"
        "Seção de Patrimônio"
    )
    cell_title_paragraph = cell_title.paragraphs[0]
    cell_title_paragraph.text = title_text

    # Set font size to 12 for the text in the second column
    for paragraph in cell_title.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(12)

    # Apply center alignment to the title cell
    cell_title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    title = document.add_paragraph()
    title.alignment = 1
    title_run_2 = title.add_run(f"{titulo}\n")
    title_run_2.font.size = Pt(14)
    title_run_2.bold = True

    entry_values = [entry.get() for entry in entries]

    # Adding sections for ORIGEM and DESTINO
    bold_paragraph = document.add_paragraph()
    bold_run = bold_paragraph.add_run(f"ORIGEM: {str(entry_values[0]).upper()}")
    bold_run.bold = True

    # Adding a new bold paragraph
    bold_paragraph = document.add_paragraph()
    bold_run = bold_paragraph.add_run(f"DESTINO: {str(entry_values[1]).upper()}\n")
    bold_run.bold = True

    # Adding a table
    table = document.add_table(rows=1, cols=3)  # Initial row for headers
    table.style = "Table Grid"  # Apply table style

    # Adding headers to the table
    headers = table.rows[0].cells
    headers[0].text = 'ITEM'
    headers[1].text = 'DESCRIÇÃO'
    headers[2].text = 'QNTD'

    # Setting preferred widths for table columns
    table.autofit = False
    table.columns[0].width = Inches(0.7)  # Adjust width as needed
    table.columns[1].width = Inches(5)  # Adjust width as needed
    table.columns[2].width = Inches(0.7)  # Adjust width as needed

    # Justify text in header cells to the middle
    for header in headers:
        for paragraph in header.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(12)  # Adjust font size if needed
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run.bold = True  # Make the text bold

    # Center the entire table within the document
    table.alignment = 1  # 0=left, 1=center, 2=right

    # Add data from Treeview to the table
    for item in tree.get_children():
        values = tree.item(item, 'values')
        row_cells = table.add_row().cells
        for index, value in enumerate(values):
            row_cells[index].text = value
            p = row_cells[index].paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_paragraph(
        "\n\n\n_____/_____/________                                    _____/_____/________                                    _____/_____/________\n").bold = True
    document.add_paragraph(
        "____________________                                      ____________________                                    ____________________").bold = True
    document.add_paragraph(
        f"   PATRIMÔNIO                                                 {assinatura1}                                                {assinatura2}").bold = True

    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])

    if file_path:
        try:
            document.save(file_path)
            os.startfile(file_path)
        except Exception as e:
            messagebox.showerror("Erro", str(e))
